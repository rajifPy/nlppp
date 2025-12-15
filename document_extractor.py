import os
import re
import logging
import tempfile
import requests
from typing import Tuple, Dict, Optional, List

# PDF Libraries
import PyPDF2
import pdfplumber
import fitz  # PyMuPDF

# DOCX Library
from docx import Document

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """
    Advanced document extractor with:
    - Text extraction from PDF, DOCX, DOC, TXT
    - Structured data extraction (Title, Abstract, Keywords, Authors, Year)
    - DOI detection and metadata fetching
    - Multi-method PDF extraction with fallback
    """
    
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx', '.doc', '.txt', '.rtf', '.md']
        self.doi_pattern = re.compile(r'\b(10\.\d{4,9}/[-._;()/:A-Z0-9]+)\b', re.IGNORECASE)
    
    # ===== MAIN EXTRACTION METHODS =====
    
    @staticmethod
    def extract_from_bytes(file_bytes: bytes, filename: str) -> Tuple[str, str, bool]:
        """
        Extract plain text from file bytes
        
        Args:
            file_bytes: Raw file content
            filename: Original filename
            
        Returns:
            Tuple[text, file_type, success]
        """
        try:
            if not filename or filename.strip() == '':
                return "", "Invalid filename", False
            
            _, ext = os.path.splitext(filename)
            ext = ext.lower()
            
            supported_ext = ['.pdf', '.docx', '.doc', '.txt', '.rtf', '.md']
            if ext not in supported_ext:
                return "", f"Unsupported format: {ext}", False
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            
            # Extract based on extension
            if ext == '.pdf':
                text = DocumentExtractor._extract_pdf_multi_method(tmp_path)
                file_type = "PDF"
            elif ext == '.docx':
                text = DocumentExtractor._extract_docx(tmp_path)
                file_type = "DOCX"
            elif ext == '.doc':
                text = DocumentExtractor._extract_doc(tmp_path)
                file_type = "DOC"
            elif ext in ['.txt', '.rtf', '.md']:
                text = DocumentExtractor._extract_text(tmp_path)
                file_type = "TEXT"
            else:
                os.unlink(tmp_path)
                return "", f"Unsupported format: {ext}", False
            
            # Clean up
            os.unlink(tmp_path)
            
            if not text or not text.strip():
                return "", "Empty text after extraction", False
            
            return text, file_type, True
            
        except Exception as e:
            logger.error(f"Extraction error for {filename}: {str(e)}")
            if 'tmp_path' in locals():
                try:
                    os.unlink(tmp_path)
                except:
                    pass
            return "", f"Extraction error: {str(e)}", False
    
    @staticmethod
    def extract_structured(file_bytes: bytes, filename: str) -> Tuple[Dict, str, bool]:
        """
        Extract document with full structure including DOI metadata
        
        Args:
            file_bytes: Raw file content
            filename: Original filename
            
        Returns:
            Tuple[structured_data, file_type, success]
            
        structured_data format:
        {
            "title": str,
            "abstract": str,
            "keywords": list[str],
            "full_text": str,
            "authors": list[str],
            "year": str,
            "doi": str,
            "publisher": str,
            "doi_metadata": dict
        }
        """
        try:
            if not filename or filename.strip() == '':
                return {}, "Invalid filename", False
            
            _, ext = os.path.splitext(filename)
            ext = ext.lower()
            
            if ext not in ['.pdf', '.docx', '.doc', '.txt']:
                return {}, f"Unsupported format for structured extraction: {ext}", False
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            
            # Extract raw text
            if ext == '.pdf':
                full_text = DocumentExtractor._extract_pdf_multi_method(tmp_path)
                file_type = "PDF"
                # Try to extract DOI from PDF
                doi = DocumentExtractor._extract_doi_from_pdf(tmp_path)
            elif ext == '.docx':
                full_text = DocumentExtractor._extract_docx(tmp_path)
                file_type = "DOCX"
                doi = DocumentExtractor._extract_doi_from_text(full_text)
            elif ext == '.doc':
                full_text = DocumentExtractor._extract_doc(tmp_path)
                file_type = "DOC"
                doi = DocumentExtractor._extract_doi_from_text(full_text)
            else:
                full_text = DocumentExtractor._extract_text(tmp_path)
                file_type = "TEXT"
                doi = DocumentExtractor._extract_doi_from_text(full_text)
            
            # Clean up temp file
            os.unlink(tmp_path)
            
            if not full_text or not full_text.strip():
                return {}, "Empty text after extraction", False
            
            # Parse structure from text
            structured = DocumentExtractor._parse_structure(full_text)
            
            # Add DOI if found
            structured["doi"] = doi if doi else ""
            
            # Fetch DOI metadata if available
            if doi:
                logger.info(f"DOI found: {doi}. Fetching metadata...")
                doi_metadata = DocumentExtractor._fetch_doi_metadata(doi)
                
                if doi_metadata and doi_metadata.get('success'):
                    # Enhance structured data with DOI metadata
                    if not structured["title"] or structured["title"] == "Untitled Document":
                        structured["title"] = doi_metadata.get('title', '')
                    
                    if not structured["abstract"]:
                        # Some APIs may provide abstract
                        structured["abstract"] = doi_metadata.get('abstract', '')[:2000]
                    
                    if not structured["authors"]:
                        structured["authors"] = doi_metadata.get('authors_list', [])
                    
                    if not structured["year"]:
                        structured["year"] = str(doi_metadata.get('year', ''))
                    
                    structured["publisher"] = doi_metadata.get('publisher', '')
                    structured["doi_metadata"] = doi_metadata
                    
                    logger.info(f"✓ DOI metadata successfully integrated")
            
            return structured, file_type, True
            
        except Exception as e:
            logger.error(f"Structured extraction error for {filename}: {str(e)}")
            if 'tmp_path' in locals():
                try:
                    os.unlink(tmp_path)
                except:
                    pass
            return {}, f"Extraction error: {str(e)}", False
    
    # ===== DOI EXTRACTION METHODS =====
    
    @staticmethod
    def _extract_doi_from_pdf(pdf_path: str) -> Optional[str]:
        """
        Smart DOI extraction from PDF with density filter
        Avoids extracting DOIs from reference sections
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            DOI string or None
        """
        try:
            doc = fitz.open(pdf_path)
            candidates = []
            
            # Step 1: Check PDF metadata
            metadata = doc.metadata
            if metadata:
                for key, value in metadata.items():
                    if value:
                        matches = DocumentExtractor._find_doi_in_text(str(value))
                        if matches:
                            logger.info("✓ DOI found in PDF metadata")
                            doc.close()
                            return matches[0]
            
            # Step 2: Scan pages with density filter
            for page_num, page in enumerate(doc):
                text = page.get_text()
                matches = DocumentExtractor._find_doi_in_text(text)
                
                # Clean matches (remove trailing punctuation)
                clean_matches = [m.rstrip('.') for m in matches]
                
                # Density filter: if too many DOIs on one page (>3), likely references section
                if len(clean_matches) > 3:
                    continue  # Skip this page
                
                # Store candidates from "clean" pages
                if clean_matches:
                    for doi in clean_matches:
                        candidates.append({
                            'doi': doi,
                            'page': page_num
                        })
            
            doc.close()
            
            # Step 3: Select best candidate
            if candidates:
                # Return earliest DOI (usually the paper's own DOI, not references)
                return candidates[0]['doi']
            
            return None
            
        except Exception as e:
            logger.error(f"DOI extraction error: {str(e)}")
            return None
    
    @staticmethod
    def _extract_doi_from_text(text: str) -> Optional[str]:
        """
        Extract DOI from plain text
        
        Args:
            text: Text content
            
        Returns:
            DOI string or None
        """
        matches = DocumentExtractor._find_doi_in_text(text)
        if matches:
            # Return first match, cleaned
            return matches[0].rstrip('.')
        return None
    
    @staticmethod
    def _find_doi_in_text(text: str) -> List[str]:
        """
        Find all DOI patterns in text
        
        Args:
            text: Text to search
            
        Returns:
            List of DOI strings
        """
        doi_pattern = re.compile(r'\b(10\.\d{4,9}/[-._;()/:A-Z0-9]+)\b', re.IGNORECASE)
        matches = doi_pattern.findall(text)
        return matches
    
    @staticmethod
    def _fetch_doi_metadata(doi: str) -> Optional[Dict]:
        """
        Fetch metadata from DOI using CrossRef/DOI.org API
        
        Args:
            doi: DOI string
            
        Returns:
            Dictionary with metadata or None
        """
        base_url = "https://doi.org/"
        url = base_url + doi
        
        # This header is CRITICAL - requests JSON citation data instead of HTML
        headers = {
            "Accept": "application/vnd.citationstyles.csl+json; charset=utf-8"
        }
        
        try:
            response = requests.get(url, headers=headers, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                
                # Extract title
                title = data.get('title', 'Title not found')
                
                # Extract year
                try:
                    published_year = data['issued']['date-parts'][0][0]
                except (KeyError, IndexError):
                    published_year = "Unknown year"
                
                # Extract authors
                authors_list = []
                if 'author' in data:
                    for author in data['author']:
                        given = author.get('given', '')
                        family = author.get('family', '')
                        full_name = f"{given} {family}".strip()
                        if full_name:
                            authors_list.append(full_name)
                
                authors_str = ", ".join(authors_list) if authors_list else "Unknown authors"
                
                return {
                    "success": True,
                    "doi": doi,
                    "title": title,
                    "year": published_year,
                    "authors": authors_str,
                    "authors_list": authors_list,
                    "publisher": data.get('publisher', 'Unknown publisher'),
                    "url": data.get('URL', url),
                    "abstract": data.get('abstract', ''),  # May not always be available
                    "type": data.get('type', 'Unknown type'),
                    "container_title": data.get('container-title', [''])[0] if data.get('container-title') else ''
                }
            
            elif response.status_code == 404:
                return {"success": False, "error": "DOI not found in database"}
            else:
                return {"success": False, "error": f"Failed to fetch. Status: {response.status_code}"}
                
        except requests.Timeout:
            return {"success": False, "error": "Request timeout"}
        except Exception as e:
            return {"success": False, "error": f"Connection error: {e}"}
    
    # ===== PDF EXTRACTION METHODS =====
    
    @staticmethod
    def _extract_pdf_multi_method(file_path: str) -> str:
        """
        Extract PDF using multiple methods with fallback
        Priority: PyMuPDF (fitz) > pdfplumber > PyPDF2
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        text = ""
        
        # Method 1: PyMuPDF (fitz) - Best for academic papers
        try:
            doc = fitz.open(file_path)
            for page in doc:
                page_text = page.get_text()
                if page_text:
                    text += page_text + "\n\n"
            doc.close()
            
            if text.strip():
                logger.info("✓ PDF extracted using PyMuPDF")
                return text.strip()
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}")
        
        # Method 2: pdfplumber - Good for tables and layout
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            
            if text.strip():
                logger.info("✓ PDF extracted using pdfplumber")
                return text.strip()
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")
        
        # Method 3: PyPDF2 - Fallback
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            
            if text.strip():
                logger.info("✓ PDF extracted using PyPDF2")
                return text.strip()
        except Exception as e:
            logger.error(f"All PDF extraction methods failed: {e}")
        
        return text.strip()
    
    # ===== STRUCTURE PARSING =====
    
    @staticmethod
    def _parse_structure(text: str) -> Dict:
        """
        Parse text to extract structured elements
        - Title
        - Abstract
        - Keywords
        - Authors
        - Year
        
        Args:
            text: Full document text
            
        Returns:
            Dictionary with structured data
        """
        lines = text.split('\n')
        
        result = {
            "title": "",
            "abstract": "",
            "keywords": [],
            "full_text": text,
            "authors": [],
            "year": "",
            "publisher": ""
        }
        
        # Clean lines
        lines = [line.strip() for line in lines if line.strip()]
        
        # ===== EXTRACT TITLE =====
        title_candidates = []
        for i, line in enumerate(lines[:15]):  # Check first 15 lines
            if 10 < len(line) < 250:
                # Prioritize UPPERCASE or Title Case
                if line.isupper() or line.istitle():
                    title_candidates.append((i, line, len(line)))
        
        if title_candidates:
            # Sort by position (earlier is better), then by length (longer is better)
            title_candidates.sort(key=lambda x: (x[0], -x[2]))
            result["title"] = title_candidates[0][1]
        else:
            # Fallback: first substantial line
            for line in lines[:5]:
                if 10 < len(line) < 250:
                    result["title"] = line
                    break
        
        # ===== EXTRACT ABSTRACT =====
        abstract_patterns = [
            r'(?i)^abstract[:\s]*(.+?)(?=\n\n|keywords|introduction|$)',
            r'(?i)^summary[:\s]*(.+?)(?=\n\n|keywords|introduction|$)',
            r'(?i)^overview[:\s]*(.+?)(?=\n\n|keywords|introduction|$)',
        ]
        
        for pattern in abstract_patterns:
            match = re.search(pattern, text, re.DOTALL | re.MULTILINE)
            if match:
                abstract_text = match.group(1).strip()
                abstract_text = re.sub(r'\s+', ' ', abstract_text)
                result["abstract"] = abstract_text[:2000]
                break
        
        # Fallback: text between title and keywords/introduction
        if not result["abstract"] and result["title"]:
            title_pos = text.lower().find(result["title"].lower())
            
            keyword_markers = ["keywords", "keyword", "key words", "introduction", "1. introduction", "1 introduction"]
            next_section_pos = len(text)
            
            for marker in keyword_markers:
                pos = text.lower().find(marker, title_pos + len(result["title"]))
                if pos != -1 and pos < next_section_pos:
                    next_section_pos = pos
            
            if title_pos < next_section_pos:
                potential_abstract = text[title_pos + len(result["title"]):next_section_pos].strip()
                potential_abstract = re.sub(r'\s+', ' ', potential_abstract)
                
                if 100 < len(potential_abstract) < 2000:
                    result["abstract"] = potential_abstract
        
        # ===== EXTRACT KEYWORDS =====
        keyword_patterns = [
            r'(?i)keywords?[:\s]*(.+?)(?=\n\n|introduction|abstract|$)',
            r'(?i)key\s*words?[:\s]*(.+?)(?=\n\n|introduction|abstract|$)',
            r'(?i)index\s*terms?[:\s]*(.+?)(?=\n\n|introduction|abstract|$)',
        ]
        
        for pattern in keyword_patterns:
            match = re.search(pattern, text, re.DOTALL | re.MULTILINE)
            if match:
                keywords_text = match.group(1).strip()
                
                # Split by common separators
                keywords = re.split(r'[;,•·\n]+', keywords_text)
                keywords = [kw.strip() for kw in keywords if kw.strip()]
                
                # Clean keywords
                cleaned_keywords = []
                for kw in keywords:
                    # Remove numbering
                    kw = re.sub(r'^\d+[\.\)]\s*', '', kw)
                    kw = re.sub(r'\s+', ' ', kw).strip()
                    
                    if 2 < len(kw) < 60:
                        cleaned_keywords.append(kw)
                
                result["keywords"] = cleaned_keywords[:20]
                break
        
        # ===== EXTRACT AUTHORS =====
        author_patterns = [
            r'(?i)^authors?[:\s]*(.+?)(?=\n\n|abstract|keywords|$)',
            r'(?i)^by[:\s]*(.+?)(?=\n\n|abstract|keywords|$)',
        ]
        
        for pattern in author_patterns:
            match = re.search(pattern, text, re.MULTILINE)
            if match:
                authors_text = match.group(1).strip()
                authors = re.split(r',|\s+and\s+|\n', authors_text)
                authors = [a.strip() for a in authors if a.strip()]
                result["authors"] = authors[:10]
                break
        
        # ===== EXTRACT YEAR =====
        year_pattern = r'\b(19|20)\d{2}\b'
        for line in lines[:25]:
            match = re.search(year_pattern, line)
            if match:
                result["year"] = match.group(0)
                break
        
        return result
    
    # ===== OTHER FORMAT EXTRACTORS =====
    
    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            return ""
    
    @staticmethod
    def _extract_doc(file_path: str) -> str:
        """Basic extraction for legacy DOC format"""
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
                
                # Try different encodings
                for encoding in ['utf-8', 'latin-1', 'windows-1252', 'cp1252']:
                    try:
                        text = content.decode(encoding, errors='ignore')
                        lines = [line.strip() for line in text.split('\n') 
                                if len(line.strip()) > 3]
                        return "\n".join(lines)
                    except:
                        continue
                
                return "Cannot extract text from .doc file"
        except Exception as e:
            logger.error(f"DOC extraction error: {str(e)}")
            return ""
    
    @staticmethod
    def _extract_text(file_path: str) -> str:
        """Extract from plain text files"""
        try:
            for encoding in ['utf-8', 'latin-1', 'windows-1252', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                        if content.strip():
                            return content
                except UnicodeDecodeError:
                    continue
            
            # Binary fallback
            with open(file_path, 'rb') as f:
                content = f.read()
                return content.decode('utf-8', errors='ignore')
                
        except Exception as e:
            logger.error(f"Text extraction error: {str(e)}")
            return ""
    
    def is_supported(self, filename: str) -> bool:
        """Check if file format is supported"""
        _, ext = os.path.splitext(filename.lower())
        return ext in self.supported_extensions


# ===== TEST CODE =====
if __name__ == "__main__":
    print("\n" + "="*80)
    print("ENHANCED DOCUMENT EXTRACTOR TEST")
    print("="*80)
    
    # Test sample text
    sample_paper = """
RENEWABLE ENERGY SYSTEMS FOR SUSTAINABLE DEVELOPMENT

John Doe¹, Jane Smith², Michael Johnson³

¹Department of Environmental Science, University of ABC
²Institute of Renewable Energy, XYZ Research Center  
³School of Engineering, DEF University

2023

DOI: 10.1016/j.renene.2023.01234

ABSTRACT

This paper explores the role of renewable energy systems in achieving sustainable 
development goals. We analyze solar, wind, and hydroelectric power generation 
technologies and their environmental impacts. Our findings demonstrate that 
transitioning to renewable energy sources can significantly reduce carbon emissions 
while promoting economic growth and energy security.

KEYWORDS: renewable energy, sustainable development, solar power, wind energy, 
carbon emissions, climate change, energy transition, green technology

1. INTRODUCTION

The global energy landscape is undergoing a fundamental transformation...
    """
    
    extractor = DocumentExtractor()
    
    print("\n📄 Testing structure parsing...")
    structured = extractor._parse_structure(sample_paper)
    
    print("\n✅ EXTRACTED STRUCTURE:")
    print(f"\n📌 Title: {structured['title']}")
    print(f"\n📝 Abstract: {structured['abstract'][:200]}...")
    print(f"\n🏷️  Keywords ({len(structured['keywords'])}): {', '.join(structured['keywords'][:5])}")
    print(f"\n👥 Authors ({len(structured['authors'])}): {', '.join(structured['authors'])}")
    print(f"\n📅 Year: {structured['year']}")
    
    # Test DOI extraction
    print("\n\n🔍 Testing DOI extraction...")
    doi = extractor._extract_doi_from_text(sample_paper)
    print(f"DOI found: {doi}")
    
    if doi:
        print(f"\n📡 Fetching DOI metadata from CrossRef...")
        metadata = extractor._fetch_doi_metadata(doi)
        
        if metadata and metadata.get('success'):
            print("\n✅ DOI METADATA:")
            print(f"Title: {metadata.get('title')}")
            print(f"Year: {metadata.get('year')}")
            print(f"Authors: {metadata.get('authors')}")
            print(f"Publisher: {metadata.get('publisher')}")
            print(f"URL: {metadata.get('url')}")
        else:
            print(f"❌ Failed to fetch metadata: {metadata.get('error')}")
    
    print("\n" + "="*80)
    print("✓ Test completed!")
    print("="*80)
